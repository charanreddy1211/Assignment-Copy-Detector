# from flask import Flask, render_template, request, jsonify, send_from_directory
# from flask_cors import CORS
# import os
# from werkzeug.utils import secure_filename
# import PyPDF2
# import docx
# import requests
# import json
# import re
# import traceback
# from difflib import SequenceMatcher

# app = Flask(__name__)
# CORS(app)

# # Configuration
# UPLOAD_FOLDER = 'uploads'
# ALLOWED_EXTENSIONS = {'pdf', 'docx'}
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# # OpenRouter API Key - Add your key here
# OPENROUTER_API_KEY = "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"  # Replace with your actual API key

# # Create uploads folder if it doesn't exist
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(file_path):
#     """Extract text from PDF file"""
#     try:
#         text = ""
#         with open(file_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             for page in pdf_reader.pages:
#                 extracted = page.extract_text()
#                 if extracted:
#                     text += extracted + "\n"
#         return text.strip()
#     except Exception as e:
#         print(f"Error extracting PDF: {e}")
#         traceback.print_exc()
#         return None

# def extract_text_from_docx(file_path):
#     """Extract text from DOCX file"""
#     try:
#         doc = docx.Document(file_path)
#         text = ""
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
#         return text.strip()
#     except Exception as e:
#         print(f"Error extracting DOCX: {e}")
#         traceback.print_exc()
#         return None

# def extract_text(file_path, file_type):
#     """Extract text based on file type"""
#     if file_type == 'pdf':
#         return extract_text_from_pdf(file_path)
#     elif file_type == 'docx':
#         return extract_text_from_docx(file_path)
#     return None

# def find_similar_lines(text1, text2, threshold=0.6):
#     """Find similar lines between two texts"""
#     lines1 = [line.strip() for line in text1.split('\n') if line.strip()]
#     lines2 = [line.strip() for line in text2.split('\n') if line.strip()]
    
#     similar_pairs = []
    
#     for i, line1 in enumerate(lines1):
#         for j, line2 in enumerate(lines2):
#             # Calculate similarity ratio
#             ratio = SequenceMatcher(None, line1.lower(), line2.lower()).ratio()
            
#             if ratio >= threshold and len(line1) > 20:  # Only consider lines with 20+ chars
#                 similar_pairs.append({
#                     'line1': line1,
#                     'line2': line2,
#                     'similarity': round(ratio * 100, 2),
#                     'line1_num': i + 1,
#                     'line2_num': j + 1
#                 })
    
#     # Sort by similarity (highest first)
#     similar_pairs.sort(key=lambda x: x['similarity'], reverse=True)
    
#     return similar_pairs[:20]  # Return top 20 matches

# def calculate_overall_similarity(text1, text2):
#     """Calculate overall similarity between two texts"""
#     # Remove extra whitespace and normalize
#     text1_clean = ' '.join(text1.split())
#     text2_clean = ' '.join(text2.split())
    
#     # Calculate similarity using SequenceMatcher
#     ratio = SequenceMatcher(None, text1_clean.lower(), text2_clean.lower()).ratio()
    
#     return round(ratio * 100, 2)

# def analyze_with_openrouter(text1, text2, api_key):
#     """Analyze code copying using OpenRouter API"""
    
#     # Limit text length to avoid token limits
#     max_chars = 3000
#     text1_truncated = text1[:max_chars] if len(text1) > max_chars else text1
#     text2_truncated = text2[:max_chars] if len(text2) > max_chars else text2
    
#     prompt = f"""You are an academic integrity expert analyzing student programming assignments for potential copying.

# Student 1 Assignment:
# {text1_truncated}

# Student 2 Assignment:
# {text2_truncated}

# Analyze these assignments for signs of copying. Consider:
# 1. Identical or very similar code structure
# 2. Same variable names and function names
# 3. Similar comments or documentation
# 4. Same logic flow and algorithms
# 5. Identical unusual solutions or approaches

# Respond ONLY with valid JSON (no markdown, no backticks):
# {{
#   "similarity_percentage": <number 0-100>,
#   "likely_copied": <boolean>,
#   "evidence": [
#     {{
#       "type": "<code_structure/variable_names/comments/logic/other>",
#       "description": "<brief explanation>",
#       "severity": "<high/medium/low>"
#     }}
#   ],
#   "conclusion": "<brief assessment of whether copying occurred>"
# }}"""

#     try:
#         print("Calling OpenRouter API...")
#         response = requests.post(
#             "https://openrouter.ai/api/v1/chat/completions",
#             headers={
#                 "Authorization": f"Bearer {api_key}",
#                 "Content-Type": "application/json",
#                 "HTTP-Referer": "http://localhost:5000",
#                 "X-Title": "Assignment Copy Checker"
#             },
#             json={
#                 "model": "anthropic/claude-3-sonnet",
#                 "messages": [
#                     {"role": "user", "content": prompt}
#                 ]
#             },
#             timeout=60
#         )
        
#         print(f"API Response Status: {response.status_code}")
        
#         if response.status_code == 200:
#             data = response.json()
            
#             if 'choices' in data and len(data['choices']) > 0:
#                 ai_response = data['choices'][0]['message']['content']
                
#                 # Clean the response
#                 cleaned_response = ai_response.strip()
#                 cleaned_response = re.sub(r'```json\n?', '', cleaned_response)
#                 cleaned_response = re.sub(r'```\n?', '', cleaned_response)
#                 cleaned_response = cleaned_response.strip()
                
#                 try:
#                     result = json.loads(cleaned_response)
#                     return result
#                 except json.JSONDecodeError as e:
#                     print(f"JSON decode error: {e}")
#                     return None
        
#         return None
            
#     except Exception as e:
#         print(f"Error in AI analysis: {e}")
#         traceback.print_exc()
#         return None

# @app.route('/favicon.ico')
# def favicon():
#     """Serve favicon or return 204 No Content"""
#     return '', 204

# @app.route('/')
# def index():
#     """Render the main page"""
#     return render_template('index.html')

# @app.route('/compare', methods=['POST'])
# def compare_assignments():
#     """Compare two uploaded assignments"""
    
#     filepath1 = None
#     filepath2 = None
    
#     try:
#         # Check if files are present
#         if 'file1' not in request.files or 'file2' not in request.files:
#             return jsonify({'error': 'Both files are required'}), 400
        
#         file1 = request.files['file1']
#         file2 = request.files['file2']
#         api_key = request.form.get('api_key', '')
        
#         print(f"Received files: {file1.filename}, {file2.filename}")
        
#         # Validate files
#         if file1.filename == '' or file2.filename == '':
#             return jsonify({'error': 'No files selected'}), 400
        
#         if not (allowed_file(file1.filename) and allowed_file(file2.filename)):
#             return jsonify({'error': 'Only PDF and DOCX files are allowed'}), 400
        
#         # Save files
#         filename1 = secure_filename(file1.filename)
#         filename2 = secure_filename(file2.filename)
#         filepath1 = os.path.join(app.config['UPLOAD_FOLDER'], filename1)
#         filepath2 = os.path.join(app.config['UPLOAD_FOLDER'], filename2)
        
#         file1.save(filepath1)
#         file2.save(filepath2)
        
#         print(f"Files saved: {filepath1}, {filepath2}")
        
#         # Extract text
#         file_type1 = filename1.rsplit('.', 1)[1].lower()
#         file_type2 = filename2.rsplit('.', 1)[1].lower()
        
#         print(f"Extracting text from {file_type1} and {file_type2}...")
        
#         text1 = extract_text(filepath1, file_type1)
#         text2 = extract_text(filepath2, file_type2)
        
#         if not text1 or not text2:
#             return jsonify({'error': 'Could not extract text from files. Please ensure files contain readable text.'}), 400
        
#         print(f"Extracted text lengths: {len(text1)}, {len(text2)}")
        
#         # Calculate similarities
#         overall_similarity = calculate_overall_similarity(text1, text2)
#         similar_lines = find_similar_lines(text1, text2, threshold=0.6)
        
#         # Get AI analysis if API key is configured
#         ai_analysis = None
#         if OPENROUTER_API_KEY and OPENROUTER_API_KEY != "your-api-key-here":
#             print("Getting AI analysis...")
#             ai_analysis = analyze_with_openrouter(text1, text2, OPENROUTER_API_KEY)
        
#         # Prepare result
#         result = {
#             'similarity_percentage': overall_similarity,
#             'likely_copied': overall_similarity > 70,
#             'similar_lines': similar_lines,
#             'total_matches': len(similar_lines),
#             'student1_file': filename1,
#             'student2_file': filename2,
#             'ai_analysis': ai_analysis
#         }
        
#         print(f"Analysis complete: Found {len(similar_lines)} similar lines")
        
#         return jsonify(result), 200
        
#     except Exception as e:
#         print(f"Error in compare_assignments: {e}")
#         traceback.print_exc()
#         return jsonify({'error': f'An error occurred: {str(e)}'}), 500
    
#     finally:
#         # Clean up files
#         try:
#             if filepath1 and os.path.exists(filepath1):
#                 os.remove(filepath1)
#                 print(f"Cleaned up: {filepath1}")
#             if filepath2 and os.path.exists(filepath2):
#                 os.remove(filepath2)
#                 print(f"Cleaned up: {filepath2}")
#         except Exception as e:
#             print(f"Error cleaning up files: {e}")

# if __name__ == '__main__':
#     print("Starting Assignment Copy Checker...")
#     print(f"Upload folder: {os.path.abspath(UPLOAD_FOLDER)}")
#     app.run(debug=True, host='0.0.0.0', port=5000)


from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import requests
import json
import re
import traceback
from difflib import SequenceMatcher
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

<<<<<<< HEAD
# OpenRouter API Key from environment
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY', '')
=======
# OpenRouter API Key - Add your key here
OPENROUTER_API_KEY = ""  # Replace with your actual API key
>>>>>>> aa26aec239ad4ffc61de039a3ce6ccdeaa80bba5

# Create uploads folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        traceback.print_exc()
        return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        traceback.print_exc()
        return None

def extract_text(file_path, file_type):
    """Extract text based on file type"""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    return None

def find_similar_lines(text1, text2, threshold=0.6):
    """Find similar lines between two texts"""
    lines1 = [line.strip() for line in text1.split('\n') if line.strip()]
    lines2 = [line.strip() for line in text2.split('\n') if line.strip()]
    
    similar_pairs = []
    
    for i, line1 in enumerate(lines1):
        for j, line2 in enumerate(lines2):
            # Calculate similarity ratio
            ratio = SequenceMatcher(None, line1.lower(), line2.lower()).ratio()
            
            if ratio >= threshold and len(line1) > 20:  # Only consider lines with 20+ chars
                similar_pairs.append({
                    'line1': line1,
                    'line2': line2,
                    'similarity': round(ratio * 100, 2),
                    'line1_num': i + 1,
                    'line2_num': j + 1
                })
    
    # Sort by similarity (highest first)
    similar_pairs.sort(key=lambda x: x['similarity'], reverse=True)
    
    return similar_pairs[:20]  # Return top 20 matches

def calculate_overall_similarity(text1, text2):
    """Calculate overall similarity between two texts"""
    # Remove extra whitespace and normalize
    text1_clean = ' '.join(text1.split())
    text2_clean = ' '.join(text2.split())
    
    # Calculate similarity using SequenceMatcher
    ratio = SequenceMatcher(None, text1_clean.lower(), text2_clean.lower()).ratio()
    
    return round(ratio * 100, 2)

def analyze_with_openrouter(text1, text2, api_key):
    """Analyze code copying using OpenRouter API"""
    
    # Limit text length to avoid token limits
    max_chars = 3000
    text1_truncated = text1[:max_chars] if len(text1) > max_chars else text1
    text2_truncated = text2[:max_chars] if len(text2) > max_chars else text2
    
    prompt = f"""You are an academic integrity expert analyzing student programming assignments for potential copying.

Student 1 Assignment:
{text1_truncated}

Student 2 Assignment:
{text2_truncated}

Analyze these assignments for signs of copying. Consider:
1. Identical or very similar code structure
2. Same variable names and function names
3. Similar comments or documentation
4. Same logic flow and algorithms
5. Identical unusual solutions or approaches

Respond ONLY with valid JSON (no markdown, no backticks):
{{
  "similarity_percentage": <number 0-100>,
  "likely_copied": <boolean>,
  "evidence": [
    {{
      "type": "<code_structure/variable_names/comments/logic/other>",
      "description": "<brief explanation>",
      "severity": "<high/medium/low>"
    }}
  ],
  "conclusion": "<brief assessment of whether copying occurred>"
}}"""

    try:
        print("Calling OpenRouter API...")
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost:5000",
                "X-Title": "Assignment Copy Checker"
            },
            json={
                "model": "anthropic/claude-3-sonnet",
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=60
        )
        
        print(f"API Response Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            
            if 'choices' in data and len(data['choices']) > 0:
                ai_response = data['choices'][0]['message']['content']
                
                # Clean the response
                cleaned_response = ai_response.strip()
                cleaned_response = re.sub(r'```json\n?', '', cleaned_response)
                cleaned_response = re.sub(r'```\n?', '', cleaned_response)
                cleaned_response = cleaned_response.strip()
                
                try:
                    result = json.loads(cleaned_response)
                    return result
                except json.JSONDecodeError as e:
                    print(f"JSON decode error: {e}")
                    return None
        
        return None
            
    except Exception as e:
        print(f"Error in AI analysis: {e}")
        traceback.print_exc()
        return None

@app.route('/favicon.ico')
def favicon():
    """Serve favicon or return 204 No Content"""
    return '', 204

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')

@app.route('/compare', methods=['POST'])
def compare_assignments():
    """Compare two uploaded assignments"""
    
    filepath1 = None
    filepath2 = None
    
    try:
        # Check if files are present
        if 'file1' not in request.files or 'file2' not in request.files:
            return jsonify({'error': 'Both files are required'}), 400
        
        file1 = request.files['file1']
        file2 = request.files['file2']
        api_key = request.form.get('api_key', '')
        
        print(f"Received files: {file1.filename}, {file2.filename}")
        
        # Validate files
        if file1.filename == '' or file2.filename == '':
            return jsonify({'error': 'No files selected'}), 400
        
        if not (allowed_file(file1.filename) and allowed_file(file2.filename)):
            return jsonify({'error': 'Only PDF and DOCX files are allowed'}), 400
        
        # Save files
        filename1 = secure_filename(file1.filename)
        filename2 = secure_filename(file2.filename)
        filepath1 = os.path.join(app.config['UPLOAD_FOLDER'], filename1)
        filepath2 = os.path.join(app.config['UPLOAD_FOLDER'], filename2)
        
        file1.save(filepath1)
        file2.save(filepath2)
        
        print(f"Files saved: {filepath1}, {filepath2}")
        
        # Extract text
        file_type1 = filename1.rsplit('.', 1)[1].lower()
        file_type2 = filename2.rsplit('.', 1)[1].lower()
        
        print(f"Extracting text from {file_type1} and {file_type2}...")
        
        text1 = extract_text(filepath1, file_type1)
        text2 = extract_text(filepath2, file_type2)
        
        if not text1 or not text2:
            return jsonify({'error': 'Could not extract text from files. Please ensure files contain readable text.'}), 400
        
        print(f"Extracted text lengths: {len(text1)}, {len(text2)}")
        
        # Calculate similarities
        overall_similarity = calculate_overall_similarity(text1, text2)
        similar_lines = find_similar_lines(text1, text2, threshold=0.6)
        
        # Get AI analysis if API key is configured
        ai_analysis = None
        if OPENROUTER_API_KEY and OPENROUTER_API_KEY != "your-api-key-here" and OPENROUTER_API_KEY.strip():
            print("Getting AI analysis...")
            ai_analysis = analyze_with_openrouter(text1, text2, OPENROUTER_API_KEY)
        
        # Prepare result
        result = {
            'similarity_percentage': overall_similarity,
            'likely_copied': overall_similarity > 70,
            'similar_lines': similar_lines,
            'total_matches': len(similar_lines),
            'student1_file': filename1,
            'student2_file': filename2,
            'ai_analysis': ai_analysis
        }
        
        print(f"Analysis complete: Found {len(similar_lines)} similar lines")
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"Error in compare_assignments: {e}")
        traceback.print_exc()
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500
    
    finally:
        # Clean up files
        try:
            if filepath1 and os.path.exists(filepath1):
                os.remove(filepath1)
                print(f"Cleaned up: {filepath1}")
            if filepath2 and os.path.exists(filepath2):
                os.remove(filepath2)
                print(f"Cleaned up: {filepath2}")
        except Exception as e:
            print(f"Error cleaning up files: {e}")

if __name__ == '__main__':
    print("Starting Assignment Copy Checker...")
    print(f"Upload folder: {os.path.abspath(UPLOAD_FOLDER)}")
<<<<<<< HEAD
    print(f"API Key configured: {bool(OPENROUTER_API_KEY and OPENROUTER_API_KEY.strip())}")
    app.run(debug=True, host='0.0.0.0', port=5000)
=======
    app.run(debug=True, host='0.0.0.0', port=5000)
>>>>>>> aa26aec239ad4ffc61de039a3ce6ccdeaa80bba5
